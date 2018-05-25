from operator import itemgetter
from tkinter import *
import tkinter as tk
from tkinter import ttk
import tkinter.messagebox
import pickle
import logging
import datetime
import tkinter.filedialog
import re
import docx
import os
from win32com import client
from docx.shared import Cm
import operator
import string
import PyPDF2
from reportlab.pdfgen.canvas import Canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer,Image,Table,TableStyle
from reportlab.pdfbase import pdfmetrics, ttfonts
pdfmetrics.registerFont(TTFont('msyh', 'STSONG.TTF'))
import os


import time

filename_Newuser = ''  # 用户信息
filename_Newreceipts = ''  # 单据信息
stageDataFile = ''  # stage信息

num=0

def mian(userfile, billsfile, stagefile):
    window = tk.Tk()
    window.title("xxx律师所")
    window.geometry('690x600')
    window.maxsize(690, 600)
    window.minsize(690, 600)
    global filename_Newuser
    global filename_Newreceipts
    global stageDataFile
    filename_Newuser = userfile  # 用户信息
    filename_Newreceipts = billsfile  # 单据信息
    stageDataFile = stagefile  # stage信息
    var_pdftitle = tk.StringVar()  # PDF标题
    var_pdfid = tk.StringVar()  # PDF编号
    var_docxtitle = tk.StringVar()# docx标题
    var_docxid = tk.StringVar()# docx标题
    #赋值序列表号
    global num
    file = open(filename_Newreceipts, 'r')
    data = file.readlines()
    if len(data)==0:
        num=0
    else:
        for dic in data:
            datadict = eval(dic)
            for i in datadict.values():
                num = i[0]
    #错误日志
    logging.basicConfig(level=logging.WARNING,
                        filename='log.txt',
                        filemode='a',
                        format='%(asctime)s - %(filename)s[line:%(lineno)d] - %(levelname)s: %(message)s')
    '''
      系统打开时将信息添加到用户列表中
    '''

    scrolly = Scrollbar(window)
    ss = []
    users = []
    userdict_user = {}
    receiptsDate = {}
    var_post = tk.StringVar()
    lbUserss = tk.Listbox(window, height=9, width=24,yscrollcommand=scrolly.set)
    scrolly.config(command=lbUserss.yview)
    fo = open(filename_Newuser, 'r')
    content = fo.readlines()
    for dic in content:
        userdict = eval(dic)
        for di in content:
            userdict1 = eval(di)
            userdict_user = userdict
        for l in userdict.keys():
            lbUserss.insert('end', l)
            users.append(l)

    fo1 = open(filename_Newreceipts, 'r')
    content1 = fo1.readlines()
    for dic in content1:
        sdict = eval(dic)
        receiptsDate = sdict

    # 添加Stage
    def addstageDate():
        try:
            if var_stageID.get()=='':
                tk.messagebox.showinfo(title='提示',message='Stage编号不能为空')
            elif check(var_stageID.get())==False:
                tk.messagebox.showinfo(title='提示',message='编号只能为数字')
            elif var_stageName.get()=='':
                tk.messagebox.showinfo(title='提示',message='Stage名称不能为空')
            elif var_stageStartDate_y.get()=='' or  var_stageStartDate_m.get()=='' or var_stage_endDate_y.get()=='' or var_stage_endDate_m.get()=='':
                tk.messagebox.showinfo(title='提示',message='日期不能为空')
            elif len(var_stageStartDate_y.get())<4 or int(var_stageStartDate_m.get())>12 or len(var_stage_endDate_y.get())<4 or int(var_stage_endDate_m.get())>12:
                tk.messagebox.showinfo(title='提示',message='日期不符合要求')
            else:
                if var_stageStartDate_y.get().isdigit()==False or var_stageStartDate_m.get().isdigit()==False or var_stage_endDate_y.get().isdigit()==False or var_stage_endDate_m.get().isdigit()==False:
                    tk.messagebox.showinfo(title='提示',message='日期只能填写数字！')
                else:
                    if int(var_stageStartDate_y.get()+('0' + var_stageStartDate_m.get() if len(
                            var_stageStartDate_m.get()) == 1 else var_stageStartDate_m.get()))< int( var_stage_endDate_y.get()+('0' + var_stage_endDate_m.get() if len(
                            var_stage_endDate_m.get())==1 else var_stage_endDate_m.get())):
                        stageData_li = []
                        file = open(stageDataFile, 'w')
                        stageID = var_stageID.get()  # stage编号
                        stageName = var_stageName.get()  # stage名称
                        stageStartDate_y = var_stageStartDate_y.get()  # stage开始时间年
                        stageStartDate_m = '0' + var_stageStartDate_m.get() if len(
                            var_stageStartDate_m.get()) == 1 else var_stageStartDate_m.get()  # stage开始时间月
                        stageEndDate_y = var_stage_endDate_y.get()  # 结束时间年
                        stageEndDate_m = '0' + var_stage_endDate_m.get() if len(
                            var_stage_endDate_m.get())==1 else var_stage_endDate_m.get()  # 结束时间月
                        stageData_li.append('Stage'+stageID)
                        stageData_li.append(stageName)
                        stageData_li.append(stageStartDate_y)
                        stageData_li.append(stageStartDate_m)
                        stageData_li.append(stageEndDate_y)
                        stageData_li.append(stageEndDate_m)
                        stageData_Z['Stage' + stageID] = stageData_li
                        file.write(str(stageData_Z))
                        file.close()
                        stage_listbox.insert('end',stageID+'  ('+stageStartDate_y+','+stageStartDate_m+'-'+stageEndDate_y+','+stageEndDate_m+')')
                        tk.messagebox.showinfo(title='提示', message='添加成功')
                        var_stageID.set('')
                        var_stageName.set('')
                        var_stageStartDate_y.set('')
                        var_stageStartDate_m.set('')
                        var_stage_endDate_y.set('')
                        var_stage_endDate_m.set('')
                    else:
                        tk.messagebox.showinfo(title='提示',message='开始日期不能大于结束日期')
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='添加失败!')
            logging.error('添加Stage日期:' + repr(e))

    # 修改stage日期
    def updateStageDate():
        try:
            if var_stageID.get()!='':
                if var_stageID.get() == '':
                    tk.messagebox.showinfo(title='提示', message='Stage编号不能为空')
                elif check(var_stageID.get()) == False:
                    tk.messagebox.showinfo(title='提示', message='编号只能为数字')
                elif var_stageName.get() == '':
                    tk.messagebox.showinfo(title='提示', message='Stage名称不能为空')
                elif var_stageStartDate_y.get() == '' or var_stageStartDate_m.get() == '' or var_stage_endDate_y.get() == '' or var_stage_endDate_m.get() == '':
                    tk.messagebox.showinfo(title='提示', message='日期不能为空')
                elif len(var_stageStartDate_y.get()) < 4 or int(var_stageStartDate_m.get()) > 12 or len(
                        var_stage_endDate_y.get()) < 4 or int(var_stage_endDate_m.get()) > 12:
                    tk.messagebox.showinfo(title='提示', message='日期不符合要求')
                else:
                    if var_stageStartDate_y.get().isdigit()==False or var_stageStartDate_m.get().isdigit()==False or var_stage_endDate_y.get().isdigit()==False or var_stage_endDate_m.get().isdigit()==False:
                        tk.messagebox.showinfo(title='提示',message='日期只能为数字')
                    else:
                        if int(var_stageStartDate_y.get() + ('0' + var_stageStartDate_m.get() if len(
                                var_stageStartDate_m.get()) == 1 else var_stageStartDate_m.get())) < int(
                            var_stage_endDate_y.get() + ('0' + var_stage_endDate_m.get() if len(
                                var_stage_endDate_m.get())==1 else var_stage_endDate_m.get())):
                            stageDate = {}
                            stageID = 'Stage'+var_stageID.get()  # stage编号
                            fo1 = open(stageDataFile, 'r')
                            content1 = fo1.readlines()
                            for dic in content1:
                                stagedict = eval(dic)
                                stagedict[stageID][1] = var_stageName.get()
                                stagedict[stageID][2] = var_stageStartDate_y.get()
                                stagedict[stageID][3] = '0' + var_stageStartDate_m.get() if len(
                                    var_stageStartDate_m.get()) == 1 else var_stageStartDate_m.get()
                                stagedict[stageID][4] = var_stage_endDate_y.get()
                                stagedict[stageID][5] = '0' + var_stage_endDate_m.get() if len(
                                    var_stage_endDate_m.get()) == 1 else var_stage_endDate_m.get()
                                stageDate = stagedict
                            fo1 = open(stageDataFile, 'w')
                            fo1.write(str(stageDate))
                            fo1.close()
                            stage_listbox.delete(0, 'end')
                            stageDate = open(stageDataFile, 'r')
                            content_stage = stageDate.readlines()
                            for i in content_stage:
                                stage = eval(i)
                                for u in stage.values():
                                    stage_listbox.insert('end',u[0][5:] + '   ' + '(' + u[2] + ',' + u[3] + '-' + u[4] + ',' +
                                                         u[5] + ')')
                            var_stageID.set('')
                            var_stageName.set('')
                            var_stageStartDate_y.set('')
                            var_stageStartDate_m.set('')
                            var_stage_endDate_y.set('')
                            var_stage_endDate_m.set('')
                            tk.messagebox.showinfo(title='提示', message='修改成功!')
                        else:
                            tk.messagebox.showinfo(title='提示',message='开始日期不能大于结束日期')
            else:
                tk.messagebox.showinfo(title='提示',message='请选择Stage,再进行修改')
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='修改失败!')
            logging.error('修改Stage日期:' + repr(e))

    # 删除stage日期
    def removeStageDate():
        try:
            if var_stageID.get()!='':
                stageID = 'Stage'+var_stageID.get()
                fo_stage = open(stageDataFile, 'r')
                content_stage = fo_stage.readlines()
                stagedict_stage = {}
                for di in content_stage:
                    userdict1 = eval(di)
                    stagedict_stage = userdict1
                    stagedict_stage.pop(stageID)
                    stageData_Z.pop(stageID)
                text = str(stagedict_stage)
                us = open(stageDataFile, 'w')
                us.write(text)
                us.close()
                stage_listbox.delete(0, 'end')
                for u in stagedict_stage.values():
                    stage_listbox.insert('end',u[0][5:] + '   ' + '(' + u[2] + ',' + u[3] + '-' + u[4] + ',' + u[5] + ')')
                var_stageID.set('')
                var_stageName.set('')
                var_stageStartDate_y.set('')
                var_stageStartDate_m.set('')
                var_stage_endDate_y.set('')
                var_stage_endDate_m.set('')
                tk.messagebox.showinfo(title='提示', message='删除成功')
            else:
                tk.messagebox.showinfo(title='提示',message='请选择Stage,再进行删除')
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='删除失败!')
            logging.error('删除Stage日期:' + repr(e))

    # 显示stage明细
    def stage1(event):
        try:
            raise_frame(stage_page)
            stageID = ('Stage'+stage_listbox.get(stage_listbox.curselection()))[: 7].strip()
            fo1 = open(stageDataFile, 'r')
            content1 = fo1.readlines()
            for dic in content1:
                stagedict = eval(dic)
                var_stageID.set(stageID[5:])
                var_stageName.set(stagedict[stageID][1])
                var_stageStartDate_y.set(stagedict[stageID][2])
                var_stageStartDate_m.set(stagedict[stageID][3])
                var_stage_endDate_y.set(stagedict[stageID][4])
                var_stage_endDate_m.set(stagedict[stageID][5])
                startDate = (stagedict[stageID][2] + stagedict[stageID][3])
                endDate = (stagedict[stageID][4] + stagedict[stageID][5])
                showData(startDate, endDate)
        except Exception as e:
            logging.error('显示Stage日期明细:' + repr(e))

    # 显示stage里面的数据
    def showData(startDate, endDate):
        try:
            data = []
            data1 = []
            stageID = stage_listbox.get(stage_listbox.curselection())
            fo1 = open(filename_Newreceipts, 'r')
            content1 = fo1.readlines()
            items = tree.get_children()
            [tree.delete(item) for item in items]
            for dic in content1:
                stagedict = eval(dic)
                for u in stagedict.values():
                    if int(startDate) <= int(u[3].replace('-', '')[:6]) <= int(endDate):
                        data.append(u)
            data.sort(key=operator.itemgetter(3), reverse=True)
            data1 = data
            for i in data1:
                tree.insert('', 0, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
        except Exception as e:
            logging.error('显示Stage里面的数据:' + repr(e))

    def user(event):
        restsUI()
        callUpdateUser()  # 将用户相关信息赋值到控件中
        selectReceipts()  # 查询该用户相关单据

    lbUserss.bind('<Button-1>', user)
    def pdfui():
        pdfWindow()
    #导出PDF文件
    def educePDF():
        try:
            story = []
            totaltime = 0.0
            totalMoney = 0.0
            story1 = []
            total=[['','',''],['Fee Earner','Total time(Hrs)','Total']]
            nums=0
            stylesheet = getSampleStyleSheet()
            styles = getSampleStyleSheet()
            normalStyle = stylesheet['Normal']

            # 打印标题编号
            story.append(Paragraph(str(var_pdftitle.get())+': '+str(var_pdfid.get()), styles['Heading1']))

            # 打印用户信息
            com = [['','','',''],['Fee Earners', 'Admitted time', 'Title', 'Hourly Rate']]
            file_user = open(filename_Newuser, 'r')
            content1 = file_user.readlines()
            for dic in content1:
                user = eval(dic)
                for i in user:
                    ss = []
                    ss.append(i+'('+str(user[i][0])+')')
                    ss.append(str(user[i][1]))
                    ss.append(str(user[i][2]))
                    ss.append('$'+str(user[i][4] * 60))
                    com.append(ss)
            component_table = Table(com, colWidths=[140, 140, 140, 140, 140])
            component_table.setStyle(TableStyle([
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('FONTNAME', (0, 0), (-1, -1), 'msyh'),
                ('GRID', (0, 2), (4, 0), 0.5, colors.black)
            ]))
            story.append(component_table)
            billsData_dict = []
            file_receipts = open(filename_Newreceipts, 'r')
            content2 = file_receipts.readlines()
            for dic in content2:
                receipts = eval(dic)
                for u in receipts.values():
                    billsData_list = []
                    billsData_list.append(u[0])
                    billsData_list.append(u[1])
                    billsData_list.append(u[2])
                    billsData_list.append(int(re.sub('-', '', u[3])))
                    billsData_list.append(u[4])
                    billsData_list.append(u[5])
                    billsData_list.append(u[6])
                    billsData_list.append(u[7])
                    billsData_list.append(u[8])
                    billsData_dict.append(billsData_list)
            billsData_dict.sort(key=operator.itemgetter(3))

            # 打印Stage信息
            stagedata = []
            stageDate = open(stageDataFile, 'r')
            content = stageDate.readlines()
            for dic in content:
                sdict = eval(dic)
                for s in sdict.values():
                    stagedata.append(s)
            suminn_time = []
            suminn_moeny = []
            for i in stagedata:
                data = []
                moeny = 0.0
                bissdata = [['', '', '', '', ''], ['Date', 'Fee Earner', 'Hours', 'Narrative', 'Total']]
                text = '<para autoLeading="off" fontSize=9><br/><br/></para>'
                story.append(Paragraph(text, normalStyle))
                story.append(
                    Paragraph(i[0] + ': ' + i[1] + '(' + i[2] + '/' + i[3] + '-' + i[4] + '/' + i[5] + ')',styles['Title'] ))
                for j in billsData_dict:
                    if int(i[2] + i[3]) <= int(str(j[3])[:6]) <= int(i[4] + i[5]):
                        ll = []
                        ll.append(str(j[3])[:4] + '-' + str(j[3])[4:6] + '-' + str(j[3])[6:8])
                        ll.append(j[1])
                        ll.append(str(round(j[4] / 60, 1)))
                        ll.append(j[2])
                        ll.append('$' + str(round(float(j[8]), 2)))
                        bissdata.append(ll)
                        data.append(ll)
                component_table11 = Table(bissdata, colWidths=[70, 70, 40, 290, 90])
                component_table11.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 12),
                    ('FONTNAME', (0, 0), (-1, -1), 'msyh'),
                    ('GRID', (0, 2), (5, 0), 0.5, colors.black)
                ]))
                story.append(component_table11)
                result_time = {}
                result_money = {}
                l = []
                for d in data:
                    result_time[d[1]] = round(float(result_time.get(d[1], 0)) + float(d[2]), 2)
                    result_money[d[1]] = round(float(result_money.get(d[1], 0)) + float(d[4][1:]), 1)
                l.append(result_time)
                l.append(result_money)
                for y in result_time:
                    data=[]
                    data.append(y)
                    data.append(result_time[y])
                    suminn_time.append(data)
                for y2 in result_money:
                    data=[]
                    data.append(y2)
                    data.append(result_money[y2])
                    suminn_moeny.append(data)
                story.append(Paragraph('———————————————————————————————', normalStyle))
                dic = {}
                for _ in l:
                    for k, v in _.items():
                        dic.setdefault(k, []).append(v)
                money_list = [ ['','',''],['Fee Earners', 'Hours', 'Total']]
                for ss in dic:
                    totaltime+=dic[ss][0]
                    totalMoney+=dic[ss][1]
                    s = []
                    s.append(ss)
                    s.append(str(dic[ss][0]))
                    s.append('$' + str(dic[ss][1]))
                    money_list.append(s)
                    moeny += dic[ss][1]
                moeny_o = []
                moeny_o.append('')
                moeny_o.append('')
                moeny_o.append('TOTAL:$' + str(moeny))
                money_list.append(moeny_o)
                li = ['', '', '']
                money_list.append(li)
                money_list1=[ ['Stage Summary:', '', '']]
                component_tablel3=Table(money_list1, colWidths=[180, 180, 180])
                component_tablel3.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 20),
                    ('FONTNAME', (0, 0), (-1, -1), 'msyh'),
                    ('FONTSIZE', (0, 0), (-1, -1), 12),  # 字体大小
                ]))

                component_tablel2 = Table(money_list, colWidths=[180, 180, 180])
                component_tablel2.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 12),
                    ('FONTNAME', (0, 0), (-1, -1), 'msyh'),
                    ('GRID', (0, 2), (3, 0), 0.5, colors.black)
                ]))
                story.append(component_tablel3)
                story.append(component_tablel2)
                nums += 1
                doc = SimpleDocTemplate('PDF/导出数据' + str(nums) + '.pdf')
                doc.build(story)
            p=[]
            result_time = dict()
            for data in suminn_time:
                result_time[data[0]] = float(result_time.get(data[0], 0)) + float(data[1])
            result_money = dict()
            for data in suminn_moeny:
                result_money[data[0]] = float(result_money.get(data[0], 0)) + float(data[1])
            sum=[]
            sum.append(result_time)
            sum.append(result_money)
            dic1 = {}
            for _ in sum:
                for k, v in _.items():
                    dic1.setdefault(k, []).append(v)
            for i in dic1:
                value=[]
                value.append(i)
                value.append(round(dic1[i][0],2))
                value.append('$'+str(dic1[i][1]))
                total.append(value)
            tota2=[]

            p.append('')
            p.append('Total: '+str(round(totaltime,2)))
            p.append('$'+str(round(totalMoney, 2)))
            tota2.append(p)
            summ=[['Summary: ','','']]
            component_tablel4 = Table(summ, colWidths=[180, 180,180])
            component_tablel3 = Table(total, colWidths=[ 180, 180,180])
            component_tablel5 = Table(tota2, colWidths=[180, 180, 180])
            component_tablel3.setStyle(TableStyle([
                        ('FONTSIZE', (0, 0), (-1, -1), 12),
                        ('FONTNAME', (0, 0), (-1, -1), 'msyh'),
                        ('GRID', (0, 2), (2, 0), 0.5, colors.black)
                    ]))
            story1.append(component_tablel4)
            story1.append(component_tablel3)
            story1.append(Paragraph('———————————————————————————————', normalStyle))
            story1.append(component_tablel5)
            doc1 = SimpleDocTemplate(os.getcwd() + '/PDF/导出数据'+str(nums+1)+'.pdf')
            doc1.build(story1)
            path = os.getcwd() + '/PDF'  # 文件夹目录
            files = os.listdir(path)  # 得到文件夹下的所有文件名称
            s = []
            for file in files:  # 遍历文件夹
                s.append(file)
            merger = PyPDF2.PdfFileMerger()
            for filename in s:
                merger.append(PyPDF2.PdfFileReader(os.getcwd() + '/PDF' + '/' + filename))
                os.remove(os.getcwd() + '/PDF' + '/' + filename)
            merger.write('导出数据PDF.pdf')
            tk.messagebox.showinfo(title='提示',message='导出成功')
        except PermissionError as pe:
            tk.messagebox.showinfo(title='提示', message='请关闭PDF/Word文件,再进行导出!')
        except FileNotFoundError:
            tk.messagebox.showinfo(title='提示',message='请导入相关文件!')
        except Exception as e:
            tk.messagebox.showinfo(title='提示', message='导出失败!')
            logging.error('导出数据PDf:' + repr(e))


    # 导出docx文件
    def docxui():
        docxWindow()
    def plot():
        try:
            totaltime = 0.0
            totalMoney = 0.0
            docxObj = docx.Document(docx=os.path.join(os.getcwd(), 'default.docx'))  # 在内存中创建Word对象
            docxObj.add_paragraph(var_docxtitle.get()+': '+ var_docxid.get(),style='TITLE')
            '''
                打印用户信息
            '''
            table_user = docxObj.add_table(rows=1, cols=4)
            table_user.columns[0].width = Cm(4)
            table_user.columns[1].width = Cm(4)
            table_user.columns[2].width = Cm(4)
            table_user.columns[3].width = Cm(4)
            hdr_cells2 = table_user.rows[0].cells
            hdr_cells2[0].text = 'Feeployee Earners '
            hdr_cells2[1].text = 'Admitted time'
            hdr_cells2[2].text = 'Title'
            hdr_cells2[3].text = 'Hourly Rate'
            file_user = open(filename_Newuser, 'r')
            content1 = file_user.readlines()
            for dic in content1:
                user = eval(dic)
                for i in user.keys():
                    row_cells = table_user.add_row().cells
                    row_cells[0].text = str(user[i][0])
                    row_cells[1].text = str(user[i][1])
                    row_cells[2].text = str(user[i][2])
                    row_cells[3].text = str(user[i][4] * 60)
            '''
                打印单据信息
            '''
            stagedata = []
            stageDate = open(stageDataFile, 'r')
            content = stageDate.readlines()
            for dic in content:
                sdict = eval(dic)
                for s in sdict.values():
                    stagedata.append(s)
            billsData_dict = []
            file_receipts = open(filename_Newreceipts, 'r')
            content2 = file_receipts.readlines()
            for dic in content2:
                receipts = eval(dic)
                for u in receipts.values():
                    billsData_list = []
                    billsData_list.append(u[0])
                    billsData_list.append(u[1])
                    billsData_list.append(u[2])
                    billsData_list.append(int(re.sub('-', '', u[3])))
                    billsData_list.append(u[4])
                    billsData_list.append(u[5])
                    billsData_list.append(u[6])
                    billsData_list.append(u[7])
                    billsData_list.append(u[8])
                    billsData_dict.append(billsData_list)
            billsData_dict.sort(key=operator.itemgetter(3))
            suminn_time = []
            suminn_moeny = []
            moeny = 0.0
            for i in stagedata:
                data = []
                docxObj.add_paragraph('')
                docxObj.add_paragraph()
                docxObj.add_paragraph(i[0] + '：' + i[1] + '(' + i[2] + '/' + i[3] + '-' + i[4] + '/' + i[5] + ')')
                docxObj.add_paragraph('——————————————————————————————————')
                table_bills = docxObj.add_table(rows=1, cols=5)
                table_bills.columns[0].width = Cm(4)
                table_bills.columns[1].width = Cm(3)
                table_bills.columns[2].width = Cm(2)
                table_bills.columns[3].width = Cm(7)

                hdr_cells2 = table_bills.rows[0].cells
                hdr_cells2[0].text = 'Date'
                hdr_cells2[1].text = 'Fee Earner'
                hdr_cells2[2].text = 'Hours'
                hdr_cells2[3].text = 'Narrative'
                hdr_cells2[4].text = 'Total '

                for j in billsData_dict:
                    if int(i[2] + i[3]) <= int(str(j[3])[:6]) <= int(i[4] + i[5]):
                        value = []
                        value.append(j[1])
                        value.append(round(j[4] / 60, 1))
                        value.append(j[8])
                        data.append(value)
                        row_cells1 = table_bills.add_row().cells
                        row_cells1[0].text = str(j[3])[:4] + '-' + str(j[3])[4:6] + '-' + str(j[3])[6:8]
                        row_cells1[1].text = j[1]
                        row_cells1[2].text = str(round(j[4] / 60, 1))
                        row_cells1[3].text = j[2]
                        row_cells1[4].text = '$'+str(round(float(j[8]),2))
                docxObj.add_paragraph('')
                docxObj.add_paragraph('Stage Summary:')
                table_summarize = docxObj.add_table(rows=1, cols=3)
                hdr_cells3 = table_summarize.rows[0].cells
                hdr_cells3[0].text = 'Fee Earners'
                hdr_cells3[1].text = 'Hours'
                hdr_cells3[2].text = 'Total'
                result_time = {}
                result_money = {}
                l = []
                for d in data:
                    result_time[d[0]] = round(float(result_time.get(d[0], 0)) + float(d[1]), 1)
                    result_money[d[0]] = round(float(result_money.get(d[0], 0)) + float(d[2]), 1)
                l.append(result_time)
                l.append(result_money)
                dic = {}
                for _ in l:
                    for k, v in _.items():
                        dic.setdefault(k, []).append(v)
                for ss in dic:
                    totaltime+=dic[ss][0]
                    totalMoney+=dic[ss][1]
                    row_cells1 = table_summarize.add_row().cells
                    row_cells1[0].text = ss
                    row_cells1[1].text = str(dic[ss][0])
                    row_cells1[2].text = '$'+str(dic[ss][1])
                    moeny += dic[ss][1]
                docxObj.add_paragraph('                                                                                                           '
                                                                                  +'Total:  '+'$' + str(moeny))
                for y in result_time:
                    data = []
                    data.append(y)
                    data.append(result_time[y])
                    suminn_time.append(data)
                for y2 in result_money:
                    data = []
                    data.append(y2)
                    data.append(result_money[y2])
                    suminn_moeny.append(data)
                result_time = dict()
                for data in suminn_time:
                    result_time[data[0]] = float(result_time.get(data[0], 0)) + float(data[1])
                result_money = dict()
                for data in suminn_moeny:
                    result_money[data[0]] = float(result_money.get(data[0], 0)) + float(data[1])
                sum = []
                sum.append(result_time)
                sum.append(result_money)
                dic1 = {}
                for _ in sum:
                    for k, v in _.items():
                        dic1.setdefault(k, []).append(v)

                docxObj.add_page_break()
                moeny = 0
            docxObj.add_paragraph('Summary:')
            docxObj.add_paragraph('')
            summarize = docxObj.add_table(rows=1, cols=3)
            hdr_cells4 = summarize.rows[0].cells
            hdr_cells4[0].text = 'Fee Earners'
            hdr_cells4[1].text = 'Total time(Hrs)'
            hdr_cells4[2].text = 'Total'
            for y in dic1:
                row_cells4 = summarize.add_row().cells
                row_cells4[0].text=y
                row_cells4[1].text =str(round(dic1[y][0],2))
                row_cells4[2].text = '$'+str(round(dic1[y][1],2))
            docxObj.add_paragraph('—————————————————————————————————————')
            docxObj.add_paragraph('Total:'+str(round(float(totaltime), 2))+'$'+str(round(float(totalMoney), 2)))

            docxObj.save('导出数据.docx')
            tk.messagebox.showinfo(title='提示',message='导出成功')
        except PermissionError as pe:
            tk.messagebox.showinfo(title='提示', message='请关闭PDF/Word文件,再进行导出!')
        except FileNotFoundError:
            tk.messagebox.showinfo(title='提示',message='请导入相关文件!')
        except Exception as e:
            tk.messagebox.showinfo(title='提示', message='导出失败!')
            logging.error('导出数据DOCX:' + repr(e))


    # 导入文件_用户
    def InputUser():
        try:
            global filename_Newuser
            filename_Newuser = tk.filedialog.askopenfilename()
            if filename_Newuser != '':
                userdict_user.clear()
                lbUserss.delete(0, 'end')
                file = open(filename_Newuser, 'r')
                text = file.readlines()
                users.clear()
                ss={}
                for dic in text:
                    userdict = eval(dic)
                    ss=userdict
                    for l in userdict:
                        lbUserss.insert('end', l)
                        users.append(l)
                user_box['values'] = users
                for va in ss:
                    l_list=[]
                    l_list.append(ss[va][0])
                    l_list.append(ss[va][1])
                    l_list.append(ss[va][2])
                    l_list.append(ss[va][3])
                    l_list.append(ss[va][4])
                    userdict_user[va]=l_list
                tk.messagebox.showinfo(title='提示', message='导入成功')
            else:
                pass
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='导入失败!')
            logging.error('导入用户文件:' + repr(e))

    # 导入文件_单据
    def inputReceipts():
        try:
            receiptsDate.clear()
            items = tree.get_children()
            [tree.delete(item) for item in items]
            global filename_Newreceipts
            filename_Newreceipts = tk.filedialog.askopenfilename()
            if filename_Newreceipts != '':
                try:
                    file_receipts = open(filename_Newreceipts, 'r')
                    cont = file_receipts.readlines()
                    for dic in cont:
                        receiptsdict = eval(dic)
                        s = 0
                        for i in receiptsdict.values():
                            tree.insert('', s, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
                    file_receipts.close()
                except Exception as e:
                    pass
                clue = tk.messagebox.showinfo('提示', '导入成功')
                global num
                file = open(filename_Newreceipts, 'r')
                data = file.readlines()
                if len(data) == 0:
                    num = 0
                else:
                    for dic in data:
                        datadict = eval(dic)
                        for i in datadict.values():
                            num = i[0]
                            sss=[]
                            sss.append(i[0])
                            sss.append(i[1])
                            sss.append(i[2])
                            sss.append(i[3])
                            sss.append(i[4])
                            sss.append(i[5])
                            sss.append(i[6])
                            sss.append(i[7])
                            sss.append(i[8])
                            receiptsDate[i[0]]=sss
                file.close()
            else:
                pass

        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='导入失败!')
            logging.error('导入单据文件:' + repr(e))

    # 查询用户明细
    def selectReceipts():
        try:
            user = lbUserss.get(lbUserss.curselection())
            fo_user = open(filename_Newreceipts, 'r')
            content_user = fo_user.readlines()
            items = tree.get_children()
            [tree.delete(item) for item in items]
            for di in content_user:
                try:
                    userdict1 = eval(di)
                    for i in userdict1.values():
                        if i[1] == user:
                            tree.insert('', 0, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
                except Exception as e:
                    pass
        except Exception as e:
            pass

    # 将值赋值到相关控件中
    def callUpdateUser():
        try:
            user = lbUserss.get(lbUserss.curselection())
            var_username.set(user)
            fo1 = open(filename_Newuser, 'r')
            content1 = fo1.readlines()
            for dic in content1:
                userdict = eval(dic)
                var_acronym.set(userdict[user][0])
                var_post.set(userdict[user][2])
                var_ReauthenticationTime.set(userdict[user][1])
                var_charge.set(round(float(userdict[user][4]) * 60, 2))
                var_sex.set(userdict[user][3])
        except Exception as e:
            pass

    # 修改用户
    def updateUser():
        try:
            if var_ReauthenticationTime.get().isdigit()==False :
                tk.messagebox.showinfo(title='提示',message='认证时间只能为数字')
            else:
                us = {}
                user = var_username.get()
                fo1 = open(filename_Newuser, 'r')
                content1 = fo1.readlines()
                for dic in content1:
                    userdict = eval(dic)
                    userdict[user][0] = var_acronym.get()
                    userdict[user][1] = var_ReauthenticationTime.get()
                    userdict[user][2] = var_post.get()
                    userdict[user][3] = var_sex.get()
                    userdict[user][4] = (float(var_charge.get()) / 60)
                    us = userdict
                fo1 = open(filename_Newuser, 'w')
                fo1.write(str(us))
                fo1.close()
                var_username.set('')
                var_acronym.set('')
                var_ReauthenticationTime.set('')
                var_post.set('')
                var_sex.set('')
                var_charge.set('')
                billsUI()
                tk.messagebox.showinfo(title='提示', message='修改成功!')
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='修改失败!')
            logging.error('修改用户:' + repr(e))

    # 删除用户
    def removeUser():
        try:
            username = var_username.get()
            userdict_user.pop(var_username.get())
            fo_user = open(filename_Newuser, 'r')
            content_user = fo_user.readlines()
            userdict_user1 = {}
            users.remove(username)
            user_box['values'] = users
            for di in content_user:
                userdict1 = eval(di)
                userdict_user1 = userdict1
                userdict_user1.pop(username)
            text = str(userdict_user1)
            us = open(filename_Newuser, 'w')
            us.write(text)
            us.close()
            lbUserss.delete(0, 'end')
            fo = open(filename_Newuser, 'r')
            content = fo.readlines()
            for dic in content:
                userdict = eval(dic)
                for l in userdict.keys():
                    lbUserss.insert('end', l)
            var_username.set('')
            var_acronym.set('')
            var_ReauthenticationTime.set('')
            var_post.set('')
            var_sex.set('')
            var_charge.set('')
            tk.messagebox.showinfo(title='提示', message='删除成功')
            billsUI()

        except  ValueError :
            pass
        except Exception as e:
            logging.error('删除用户:' + repr(e))
            tk.messagebox.showinfo(title='错误', message='删除失败')

    # 增加用户页面_清空按钮
    def cancels():
        var_username_add.set('')
        var_acronym_add.set('')
        var_ReauthenticationTime_add.set('')
        var_post_add.set('')
        var_sex_add.set('0')
        var_charge_add.set('')
    # 增加用户
    def addUser():
        try:
            user_list = []
            username = var_username_add.get()
            acronym = var_acronym_add.get()
            ReauthenticationTime = var_ReauthenticationTime_add.get()
            post = var_post_add.get()
            sex = '男' if var_sex_add.get() == 1 else '女'
            charge = var_charge_add.get()
            if (username == '') | (acronym == '')  | (sex == '') | (charge == ''):
                tk.messagebox.showinfo(title='提示', message='请添加相关信息!')
            elif len(ReauthenticationTime)<4:
                tk.messagebox.showinfo(title='提示',message='认证时间格式错误')
            else:
                if charge.isdigit() == False :
                    tk.messagebox.showinfo(title='错误', message='收费标准只能为数字')
                else:
                    users.append(username)
                    user_box['values'] = users
                    user_list.append(acronym)
                    user_list.append(ReauthenticationTime)
                    user_list.append(post)
                    user_list.append(sex)
                    user_list.append((int(charge) / 60))
                    userdict_user[username] = user_list

                    text = str(userdict_user)
                    us = open(filename_Newuser, 'w')
                    us.write(text)
                    us.close()
                    lbUserss.insert('end', username)
                    var_ReauthenticationTime_add.set('')
                    var_post_add.set('')
                    var_charge_add.set('')
                    var_acronym_add.set('')
                    var_sex_add.set(0)
                    tk.messagebox.showinfo(title='提示', message='添加成功!')
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='添加失败!')
            logging.error('添加用户:' + repr(e))

    # 退出系统
    def exita():
        sys.exit()


    # 添加_收据单
    def confirms():
        try:
            user = var_user.get()
            incident = var_incident.get()
            jobDate_y = var_jobDate_y.get()
            jobDate_m = '0' + var_jobDate_m.get() if len(var_jobDate_m.get()) == 1 else var_jobDate_m.get()
            jobDate_d = '0' + var_jobDate_d.get() if len(var_jobDate_d.get()) == 1 else var_jobDate_d.get()
            serDateHrs = var_serDate_hrs.get()
            serDateMins = var_serDate_mins.get()
            copying = var_copying.get()
            filing = var_filing.get()
            serving = var_serving.get()
            if user == '':
                tk.messagebox.showinfo(title='错误', message='用户不能为空!')
            elif incident == '':
                tk.messagebox.showinfo(title='错误', message='事件名称不能为空!')
            elif jobDate_y == ''or jobDate_m == '' or  jobDate_d == '':
                tk.messagebox.showinfo(title='错误', message='工作日期不能为空!')
            else:
                if jobDate_y.isdigit()==False or jobDate_m.isdigit()==False or jobDate_d.isdigit()==False  :
                    tk.messagebox.showinfo(title='提示', message='工作时间只能为数字')
                elif serDateHrs.isdigit()==False or serDateMins.isdigit()==False:
                    tk.messagebox.showinfo(title='提示', message='服务时间只能为数字')
                elif copying.isdigit() == False:
                    tk.messagebox.showinfo(title='提示', message='Copying只能为数字')
                elif filing.isdigit() == False:
                    tk.messagebox.showinfo(title='提示', message='Filing只能为数字')
                elif serving.isdigit() == False:
                    tk.messagebox.showinfo(title='提示', message='Serving只能为数字')
                elif int(jobDate_m)>12 or int(jobDate_d)>31 or len(jobDate_y)<4:
                    tk.messagebox.showinfo(title='提示',message='工作日期不符合要求')
                else:
                    charge = ''
                    time=0
                    receipts_lits = []
                    fo1 = open(filename_Newuser, 'r')
                    content1 = fo1.readlines()
                    for dic in content1:
                        userdict = eval(dic)
                        charge = userdict[user][4]
                        time= userdict[user][1]
                    if jobDate_y<time:
                        tk.messagebox.showinfo(title='提示',message='工作日期不能小于该用户的认证日期')
                    else:
                        global num
                        num+=1
                        receipts_lits.append(num)
                        receipts_lits.append(user)
                        receipts_lits.append(incident)
                        receipts_lits.append(jobDate_y + '-' + jobDate_m + '-' + jobDate_d)
                        receipts_lits.append((int(serDateHrs) * 60) + int(serDateMins))
                        receipts_lits.append(int(copying))
                        receipts_lits.append(int(filing))
                        receipts_lits.append(int(serving))
                        receipts_lits.append(
                            (float(charge) * ((int(serDateHrs) * 60) + int(serDateMins))) + int(copying) + int(
                                filing) + int(serving))
                        receiptsDate[num] = receipts_lits
                        receipts1 = open(filename_Newreceipts, 'w')
                        receipts1.write(str(receiptsDate))
                        receipts1.close()
                        tk.messagebox.showinfo(title='提示', message='添加成功!')
                        var_incident.set('')
                        var_jobDate_y.set('')
                        var_jobDate_m.set('')
                        var_jobDate_d.set('')
                        var_serDate_hrs.set('0')
                        var_serDate_mins.set('0')
                        var_copying.set('0')
                        var_filing.set('0')
                        var_serving.set('0')
                        s = 0
                        tree.insert('', s, values=(num, user,
                                                   jobDate_y + '-' + jobDate_m + '-' + jobDate_d,
                                                   (int(serDateHrs) * 60) + int(serDateMins),incident,
                                                   int(copying)+int(filing)+int(serving),
                                                   (float(charge) * ((int(serDateHrs) * 60) + int(serDateMins))) + float(
                                                       copying) + float(filing) + float(serving))
                                    )
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='添加失败!')
            logging.error('添加收据单:' + repr(e))
    #显示全部单据
    def overallData():
        try:
            items = tree.get_children()
            [tree.delete(item) for item in items]
            file_receipts = open(filename_Newreceipts, 'r')
            cont = file_receipts.readlines()
            for dic in cont:
                receiptsdict = eval(dic)
                s = 0
                for i in receiptsdict.values():
                    tree.insert('', s, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
        except Exception as e:
            pass
    def cancel():
        var_incident.set('')
        var_jobDate_y.set('0')
        var_jobDate_m.set('0')
        var_jobDate_d.set('0')
        var_serDate_hrs.set('0')
        var_serDate_mins.set('0')
        var_copying.set('0')
        var_filing.set('0')
        var_serving.set('0')

    var_incident = tk.StringVar()
    var_jobDate_y = tk.StringVar()
    var_jobDate_m = tk.StringVar()
    var_jobDate_d = tk.StringVar()
    var_serDate_hrs = tk.StringVar()
    var_serDate_mins = tk.StringVar()
    var_copying = tk.StringVar()
    var_copying.set('0')
    var_filing = tk.StringVar()
    var_filing.set('0')
    var_copying.set('0')
    var_serving = tk.StringVar()
    var_serving.set('0')
    var_user = tk.StringVar()
    var_username = tk.StringVar()
    var_acronym = tk.StringVar()
    var_ReauthenticationTime = tk.StringVar()
    var_charge = tk.StringVar()
    var_sex = tk.StringVar()
    var_serDate_hrs.set('0')
    var_serDate_mins.set('0')

    var_stageStartTime = tk.StringVar()
    var_stageEdnTime = tk.StringVar()
    # 导航条
    men = tk.Menu(window)
    usermenu = tk.Menu(men, tearoff=0)
    men.add_cascade(label='功能', menu=usermenu)
    usermenu.add_command(label='导入用户数据', command=InputUser)
    usermenu.add_command(label='导入单据数据', command=inputReceipts)
    usermenu.add_command(label='导出DOCX', command=docxui)
    usermenu.add_command(label='导出PDF', command=pdfui)
    usermenu.add_command(label='显示全部数据', command=overallData)
    exitemenu = tk.Menu(men, tearoff=0)
    men.add_cascade(label='退出', menu=exitemenu)
    exitemenu.add_command(label='退出', command=exita)
    window.config(menu=men)


    # 删除单据
    def removeBills():
        try:
            date = {}
            date_list = []
            date_list.append(var_user.get())
            date_list.append(var_incident.get())
            date_list.append(var_jobDate_y.get() + '-' + var_jobDate_m.get() + '-' + var_jobDate_d.get())
            date_list.append((int(var_serDate_hrs.get()) * 60) + int(var_serDate_mins.get()))
            fo = open(filename_Newreceipts, 'r')
            content = fo.readlines()
            for dic in content:
                sdict = eval(dic)
                date = sdict
            date.pop(int(var_serialNum.get()))
            receiptsDate.pop(int(var_serialNum.get()))
            wri = open(filename_Newreceipts, 'w')
            wri.write(str(date))
            wri.close()
            tk.messagebox.showinfo(title='提示', message='删除成功!')
            var_serialNum.set('')
            var_user.set('')
            var_incident.set('')
            var_jobDate_y.set('')
            var_jobDate_m.set('')
            var_jobDate_d.set('')
            var_serDate_hrs.set('0')
            var_serDate_mins.set('0')
            var_copying.set('0')
            var_filing.set('0')
            var_serving.set('0')
            removeBills.place_forget()
            updateBills.place_forget()
            confirm.place(x=160, y=270)
            cancel_i.place(x=230, y=270)

            items = tree.get_children()
            [tree.delete(item) for item in items]
            try:
                file_receipts = open(filename_Newreceipts, 'r')
                cont = file_receipts.readlines()
                for dic in cont:
                    receiptsdict = eval(dic)
                    s = 0
                    for i in receiptsdict.values():
                        tree.insert('', s, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
            except Exception as e:
                pass
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='删除失败')
            logging.error('删除单据:' + repr(e))

    # 修改单据
    def updateBills():
        try:
            if var_user.get() == '':
                tk.messagebox.showinfo(title='错误', message='用户不能为空!')
            elif var_incident.get() == '':
                tk.messagebox.showinfo(title='错误', message='事件名称不能为空!')
            elif var_jobDate_y.get() == '' or var_jobDate_m.get() == '' or var_jobDate_d.get() == '':
                tk.messagebox.showinfo(title='错误', message='工作日期不能为空!')
            else:
                if var_jobDate_y.get().isdigit()==False or var_jobDate_m.get().isdigit()==False or var_jobDate_d.get().isdigit()==False:
                    tk.messagebox.showinfo(title='提示',message='工作时间只能为数字')
                elif var_serDate_hrs.get().isdigit()==False or var_serDate_mins.get().isdigit()==False:
                    tk.messagebox.showinfo(title='提示',message='服务时间只能为数字')
                elif var_copying.get().isdigit()==False :
                    tk.messagebox.showinfo(title='提示',message='Copying只能为数字')
                elif var_filing.get().isdigit()==False:
                    tk.messagebox.showinfo(title='提示',message='Filing只能为数字')
                elif var_serving.get().isdigit()==False:
                    tk.messagebox.showinfo(title='提示',message='Serving只能为数字')
                else:
                    jobDate_m = '0' + var_jobDate_m.get() if len(var_jobDate_m.get()) == 1 else var_jobDate_m.get()
                    jobDate_d = '0' + var_jobDate_d.get() if len(var_jobDate_d.get()) == 1 else var_jobDate_d.get()
                    charge = ''
                    fo1 = open(filename_Newuser, 'r')
                    content1 = fo1.readlines()
                    for dic in content1:
                        userdict = eval(dic)
                        charge = userdict[var_user.get()][4]
                    date = {}
                    fo = open(filename_Newreceipts, 'r')
                    content = fo.readlines()
                    for dic in content:
                        sdict = eval(dic)
                        sdict[int(var_serialNum.get())][1] = var_user.get()
                        sdict[int(var_serialNum.get())][2] = var_incident.get()
                        sdict[int(var_serialNum.get())][3] = (var_jobDate_y.get() + '-' + jobDate_m + '-' + jobDate_d)
                        sdict[int(var_serialNum.get())][4] = (int(var_serDate_hrs.get()) * 60) + int(var_serDate_mins.get())
                        sdict[int(var_serialNum.get())][5] = var_copying.get()
                        sdict[int(var_serialNum.get())][6] = var_filing.get()
                        sdict[int(var_serialNum.get())][7] = var_serving.get()
                        sdict[int(var_serialNum.get())][8] = (
                                (float(charge) * ((int(var_serDate_hrs.get()) * 60) + int(var_serDate_mins.get()))) + int(
                            var_copying.get()) + int(var_filing.get()) + int(var_serving.get()))
                        date = sdict
                        wri = open(filename_Newreceipts, 'w')
                        wri.write(str(date))
                        wri.close()
                        var_serialNum.set('')
                        var_user.set('')
                        var_incident.set('')
                        var_jobDate_y.set('')
                        var_jobDate_m.set('')
                        var_jobDate_d.set('')
                        var_serDate_hrs.set('0')
                        var_serDate_mins.set('0')
                        var_copying.set('0')
                        var_filing.set('0')
                        var_serving.set('0')
                        removeBills.place_forget()
                        updateBills.place_forget()
                        confirm.place(x=160, y=270)
                        cancel_i.place(x=230, y=270)
                        tk.messagebox.showinfo(title='提示', message='修改成功!')
                        items = tree.get_children()
                        [tree.delete(item) for item in items]
                        try:
                            file_receipts = open(filename_Newreceipts, 'r')
                            cont = file_receipts.readlines()
                            for dic in cont:
                                receiptsdict = eval(dic)
                                s = 0
                                for i in receiptsdict.values():
                                    tree.insert('', s, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
                        except Exception as e:
                            pass
        except Exception as e:
            tk.messagebox.showinfo(title='错误', message='修改失败')
            logging.error('修改单据:' + repr(e))

    var_serialNum = tk.StringVar()  # 单据序列号
    rows = 0
    while rows < 50:
        window.rowconfigure(rows, weight=1)
        window.columnconfigure(rows, weight=1)
        rows += 1

    def raise_frame(frame):
        frame.tkraise()
    def billsUI():
        raise_frame(bills_page)
        var_stageStartDate_y.set('')
        var_stageStartDate_m.set('')
        var_stageName.set('')
        var_stageID.set('')
        var_stage_endDate_y.set('')
        var_stage_endDate_m.set('')

        var_username_add.set('')
        var_acronym_add .set('')
        var_sex_add.set(0)
        var_post_add.set('')
        var_ReauthenticationTime_add.set('')
        var_charge_add.set('')
    def userUI():
        raise_frame(adduser_page)
        var_incident.set('')
        var_jobDate_y.set('0')
        var_jobDate_m.set('0')
        var_jobDate_d.set('0')
        var_serDate_hrs.set('0')
        var_serDate_mins.set('0')
        var_copying.set('0')
        var_filing.set('0')
        var_serving.set('0')
        var_stageStartDate_y.set('')
        var_stageStartDate_m.set('')
        var_stageName.set('')
        var_stageID.set('')
        var_stage_endDate_y.set('')
        var_stage_endDate_m.set('')
        var_username_add.set('')
        var_acronym_add.set('')
        var_sex_add.set(0)
        var_post_add.set('')
        var_ReauthenticationTime_add.set('')
        var_charge_add.set('')
    def stageUI():
        var_incident.set('')
        var_jobDate_y.set('0')
        var_jobDate_m.set('0')
        var_jobDate_d.set('0')
        var_serDate_hrs.set('0')
        var_serDate_mins.set('0')
        var_copying.set('0')
        var_filing.set('0')
        var_serving.set('0')
        raise_frame(stage_page)
        var_username_add.set('')
        var_acronym_add.set('')
        var_sex_add.set(0)
        var_post_add.set('')
        var_ReauthenticationTime_add.set('')
        var_charge_add.set('')
    def restsUI():
        var_username_add.set('')
        var_acronym_add.set('')
        var_sex_add.set(0)
        var_post_add.set('')
        var_ReauthenticationTime_add.set('')
        var_charge_add.set('')
        var_incident.set('')
        var_jobDate_y.set('0')
        var_jobDate_m.set('0')
        var_jobDate_d.set('0')
        var_serDate_hrs.set('0')
        var_serDate_mins.set('0')
        var_copying.set('0')
        var_filing.set('0')
        var_serving.set('0')
        var_stageStartDate_y.set('')
        var_stageStartDate_m.set('')
        var_stageName.set('')
        var_stageID.set('')
        var_stage_endDate_y.set('')
        var_stage_endDate_m.set('')

        raise_frame(rests_page)

    bills_page = Frame(window, width=500, height=300)
    adduser_page = Frame(window, width=500, height=300)
    stage_page = Frame(window, width=500, height=300)
    rests_page=Frame(window, width=500, height=300)

    billsUI1 = tk.Button(bills_page, text='Event', width=8, command=billsUI, borderwidth=1).place(x=5,y=0)
    userUI1  = tk.Button(bills_page, text='FeeEarner', width=8, command=userUI, borderwidth=1).place(x=80,y=0)
    stageUI1 = tk.Button(bills_page, text='Stage', width=8, command=stageUI, borderwidth=1).place(x=155,y=0)

    billsUI2 = tk.Button(adduser_page, text='Event', width=8, command=billsUI, borderwidth=1).place(x=5,y=0)
    userUI2 = tk.Button(adduser_page, text='FeeEarner', width=8, command=userUI, borderwidth=1).place(x=80,y=0)
    stageUI2 = tk.Button(adduser_page, text='Stage', width=8, command=stageUI, borderwidth=1).place(x=155,y=0)

    billsUI3 = tk.Button(stage_page, text='Event', width=8, command=billsUI, borderwidth=1).place(x=5,y=0)
    userUI3  = tk.Button(stage_page, text='FeeEarner', width=8, command=userUI, borderwidth=1).place(x=80,y=0)
    stageUI3 = tk.Button(stage_page, text='Stage', width=8, command=stageUI, borderwidth=1).place(x=155,y=0)

    billsUI3 = tk.Button(rests_page, text='Event', width=8, command=billsUI, borderwidth=1).place(x=5, y=0)
    userUI3 = tk.Button(rests_page, text='FeeEarner', width=8, command=userUI, borderwidth=1).place(x=80,y=0)
    stageUI3 = tk.Button(rests_page, text='Stage', width=8, command=stageUI, borderwidth=1).place(x=155, y=0)

    for frame in (bills_page, adduser_page, stage_page,rests_page):
        frame.grid(row=0, column=0, sticky='news')



    # 添加收据单_控件
    serialNum = tk.Entry(bills_page, textvariable=var_serialNum)

    user_la = tk.Label(bills_page, text="Fee Earners:")
    user_la.place(x=50, y=40)
    user_box = ttk.Combobox(bills_page, width=12, textvariable=var_user)
    user_box['values'] = users
    user_box['state']='readonly'
    user_box.place(x=150, y=40)

    incident_la = tk.Label(bills_page, text="Narrative:")
    incident_la.place(x=50, y=70)
    incident_entry = tk.Entry(bills_page, textvariable=var_incident)
    incident_entry.place(x=150, y=70)

    startDate_la = tk.Label(bills_page, text='Date:')
    startDate_la.place(x=50, y=100)
    startDate_la1=tk.Label(bills_page,text='Y:')
    startDate_la1.place(x=150,y=100)
    startDate_la2=tk.Label(bills_page,text='M:')
    startDate_la2.place(x=210,y=100)
    startDate_la3=tk.Label(bills_page,text='D:')
    startDate_la3.place(x=270,y=100)
    startDate_entry1 = tk.Entry(bills_page, width=5, textvariable=var_jobDate_y)
    startDate_entry1.place(x=170, y=100)
    startDate_entry2 = tk.Entry(bills_page, width=5, textvariable=var_jobDate_m)
    startDate_entry2.place(x=230, y=100)
    startDate_entry3 = tk.Entry(bills_page, width=5, textvariable=var_jobDate_d)
    startDate_entry3.place(x=290, y=100)

    serDate_la2 = tk.Label(bills_page, text='Working Hours:')
    serDate_la2.place(x=50, y=130)
    serDate_la3 = tk.Label(bills_page, text='Hsr:')
    serDate_la3.place(x=150, y=130)
    serDate_entry1 = tk.Entry(bills_page, width=5, textvariable=var_serDate_hrs)
    serDate_entry1.place(x=180, y=130)
    serDate_la4 = tk.Label(bills_page, text='Mins:')
    serDate_la4.place(x=220, y=130)
    serDate_entry2 = tk.Entry(bills_page, width=5, textvariable=var_serDate_mins)
    serDate_entry2.place(x=260, y=130)

    copying_la = tk.Label(bills_page, text='Copying:')
    copying_la.place(x=50, y=160)
    copying_entry = tk.Entry(bills_page, textvariable=var_copying)
    copying_entry.place(x=150, y=160)

    filing_la = tk.Label(bills_page, text='Filing:')
    filing_la.place(x=50, y=190)
    filing_entry = tk.Entry(bills_page, textvariable=var_filing)
    filing_entry.place(x=150, y=190)

    serving_la = tk.Label(bills_page, text='Serving:')
    serving_la.place(x=50, y=220)
    serving_entry = tk.Entry(bills_page, textvariable=var_serving)
    serving_entry.place(x=150, y=220)

    confirm = tk.Button(bills_page, text='添 加', width=6, command=confirms)
    confirm.place(x=160, y=270)
    removeBills = tk.Button(bills_page, text='删 除', width=6, command=removeBills)
    updateBills = tk.Button(bills_page, text='修 改', width=6, command=updateBills)
    cancel_i = tk.Button(bills_page, text='清 空', width=6, command=cancel)
    cancel_i.place(x=230, y=270)

    # 排序_升序
    def call_back(event):
        try:
            if 116 >= event.x >= 66 and 23 >= event.y >= 1:  # 用户排序
                callBack(1)
            elif 536 >= event.x >= 312 and 23 >= event.y >= 1:  # 事件名称
                callBack(4)
            elif 207 >= event.x >= 126 and 23 >= event.y >= 1:  # 工作日期
                callBack(2)
            elif 299 >= event.x >= 218 and 23 >= event.y >= 1:  # 服务时间
                callBack(3)
            elif 617 >= event.x >= 546 and 23 >= event.y >= 1:  # 其他费用
                callBack(5)
            elif 686 >= event.x >= 628 and 23 >= event.y >= 1:  # 其他费用
                callBack(6)
            else:
                pass
        except Exception as e:
            logging.error('排序_升序:' + repr(e))

    # 排序_降序
    def callback_order(event):
        try:
            if 116 >= event.x >= 66 and 23 >= event.y >= 1:  # 用户排序
                callBack_order(1)
            elif 536 >= event.x >= 312 and 23 >= event.y >= 1:  # 事件名称
                callBack_order(4)
            elif 207 >= event.x >= 126 and 23 >= event.y >= 1:  # 工作日期
                callBack_order(2)
            elif 299 >= event.x >= 218 and 23 >= event.y >= 1:  # 服务时间
                callBack_order(3)
            elif 617 >= event.x >= 546 and 23 >= event.y >= 1:  # 合计
                callBack_order(5)
            elif 686 >= event.x >= 628 and 23 >= event.y >= 1:  # 合计
                callBack_order(6)
            else:
                pass
        except Exception as e:
            logging.error('排序_降序:' + repr(e))

    # 控制台
    tree = ttk.Treeview(window, show="headings", height=13)

    tree["columns"] = ('ID', 'Fee Earners', 'Date', 'Billable MIns', 'Title', 'Rests', 'Total')
    tree.column('ID', width=60, anchor="center")
    tree.column('Fee Earners', width=80, anchor="center")
    tree.column('Date', width=92, anchor="center")
    tree.column('Billable MIns', width=92, anchor="center")
    tree.column('Title', width=210, anchor="center")
    tree.column('Rests', width=82, anchor="center")
    tree.column('Total', width=82, anchor="center")
    tree.bind("<Button-1>", call_back)
    tree.bind("<Double-Button-1>", callback_order)
    text = []

    def trefun(event):
        billsUI()
        iids = tree.selection()
        for i in iids:
            text.append(tree.item(i, 'values'))
        for o in text:
            var_serialNum.set(o[0])
            var_user.set(o[1])
            var_incident.set(o[4])
            var_jobDate_y.set(o[2].replace('-', '')[:4])
            var_jobDate_m.set(o[2].replace('-', '')[4:6])
            var_jobDate_d.set(o[2].replace('-', '')[6:8])
            if float(o[3])<60:
                var_serDate_mins.set(o[3])
            else:
                var_serDate_hrs.set(int(float(o[3])/60))
                var_serDate_mins.set(int(float(o[3]))%60)
            confirm.place(x=160, y=260)
            updateBills.place(x=280, y=260)
            removeBills.place(x=220, y=260)
            cancel_i.place(x=340, y=260)

    tree.bind("<<TreeviewSelect>>", trefun)
    tree.column('ID', width=60, anchor="center")
    tree.column('Fee Earners', width=80, anchor="center")
    tree.column('Date', width=92, anchor="center")
    tree.column('Billable MIns', width=92, anchor="center")
    tree.column('Title', width=210, anchor="center")
    tree.column('Rests', width=82, anchor="center")
    tree.column('Total', width=82, anchor="center")


    tree.heading('ID', text='ID')
    tree.heading('Fee Earners', text='Fee Earners')
    tree.heading('Date', text='Date')
    tree.heading('Billable MIns', text='Billable MIns')
    tree.heading('Title', text='Title')
    tree.heading('Rests', text='Rests')
    tree.heading('Total', text='Total')
    tree.place(x=0, y=315)
    vbar = ttk.Scrollbar(window, orient=VERTICAL, command=tree.yview)

    # 显示用户详细UI
    usernmae_la2 = tk.Label(rests_page, text="Fee Earners:")
    username_entry2 = tk.Entry(rests_page, textvariable=var_username, state='disabled')
    acronym_la2 = tk.Label(rests_page, text="Initials:")
    acronym_entry2 = tk.Entry(rests_page, textvariable=var_acronym)
    sex_la2 = tk.Label(rests_page, text='Sex:')
    sex_man2 = tk.Radiobutton(rests_page, text='男', value='男', variable=var_sex)
    sex_woman2 = tk.Radiobutton(rests_page, text='女', value='女', variable=var_sex)
    post_la2 = tk.Label(rests_page, text='Title:')
    numberChosen2 = tk.Entry(rests_page, textvariable=var_post)
    ReauthenticationTime_la2 = tk.Label(rests_page, text='Admitted Time:')
    ReauthenticationTime_entry2 = tk.Entry(rests_page, textvariable=var_ReauthenticationTime)
    charge_la2 = tk.Label(rests_page, text='Hourly Rate:')
    charge_entry2 = tk.Entry(rests_page, textvariable=var_charge)
    updateUser = tk.Button(rests_page, text='修改', width=5, command=updateUser)
    removeUser = tk.Button(rests_page, text='删除', width=5, command=removeUser)

    usernmae_la2.place(x=50, y=35)
    username_entry2.place(x=145, y=35)
    acronym_la2.place(x=50, y=65)
    acronym_entry2.place(x=145, y=65)
    sex_la2.place(x=50, y=95)
    sex_man2.place(x=145, y=95)
    sex_woman2.place(x=180, y=95)
    post_la2.place(x=50, y=125)
    numberChosen2.place(x=145, y=125)
    ReauthenticationTime_la2.place(x=50, y=155)
    ReauthenticationTime_entry2.place(x=145, y=155)
    charge_la2.place(x=50, y=185)
    charge_entry2.place(x=145, y=185)
    removeUser.place(x=160, y=220)
    updateUser.place(x=240, y=220)
    '''
    页面加载完成时将单据添加到控制台
    '''
    # 显示添加的收据单

    try:
        file_receipts = open(filename_Newreceipts, 'r')
        cont = file_receipts.readlines()
        for dic in cont:
            receiptsdict = eval(dic)
            s = 0
            for i in receiptsdict.values():
                tree.insert('', s, values=(i[0], i[1], i[3], i[4], i[2], int(i[5])+int(i[6])+int(i[7]), round(i[8], 3)))
    except Exception as e:
        pass

    lbUserss.place(x=530, y=0)

    var_username_add = tk.StringVar()
    var_acronym_add = tk.StringVar()
    var_sex_add = tk.IntVar()
    var_post_add = tk.StringVar()
    var_ReauthenticationTime_add = tk.StringVar()
    var_charge_add = tk.StringVar()

    # 增加用户ui
    username_la = tk.Label(adduser_page, text="Fee Earners:")
    username_la.place(x=50, y=40)
    username_entry = tk.Entry(adduser_page, textvariable=var_username_add)
    username_entry.place(x=130, y=40)
    acronym_la = tk.Label(adduser_page, text="Initials:")
    acronym_la.place(x=50, y=70)
    acronym_entry = tk.Entry(adduser_page, textvariable=var_acronym_add)
    acronym_entry.place(x=130, y=70)
    sex_la = tk.Label(adduser_page, text='Sex:')
    sex_la.place(x=50, y=100)
    sex_man = tk.Radiobutton(adduser_page, text='男', value=1, variable=var_sex_add)
    sex_man.place(x=130, y=100)
    sex_woman = tk.Radiobutton(adduser_page, text='女', value=2, variable=var_sex_add)
    sex_woman.place(x=180, y=100)
    post_la = tk.Label(adduser_page, text='Title:')
    post_la.place(x=50, y=130)
    numberChosen = tk.Entry(adduser_page, textvariable=var_post_add)
    numberChosen.place(x=130, y=130)
    ReauthenticationTime_la = tk.Label(adduser_page, text='Admitted Time:')
    ReauthenticationTime_la.place(x=50, y=160)
    ReauthenticationTime_entry = tk.Entry(adduser_page, textvariable=var_ReauthenticationTime_add)
    ReauthenticationTime_entry.place(x=145, y=160)
    charge_la = tk.Label(adduser_page, text='Hourly Rate:')
    charge_la.place(x=50, y=190)
    charge_entry = tk.Entry(adduser_page, textvariable=var_charge_add)
    charge_entry.place(x=130, y=190)
    addUser = tk.Button(adduser_page, text='添 加', width=5, command=addUser)
    addUser.place(x=160, y=240)
    cancel = tk.Button(adduser_page, text='清 空', width=5, command=cancels)
    cancel.place(x=220, y=240)

    var_stageStartDate_y = tk.StringVar()
    var_stageStartDate_m = tk.StringVar()
    var_stageName = tk.StringVar()
    var_stageID = tk.StringVar()
    var_stage_endDate_y = tk.StringVar()
    var_stage_endDate_m = tk.StringVar()

    # stage_page
    stageID_La = tk.Label(stage_page, text='Stage ID:')
    stageID_La.place(x=50, y=40)
    stageID_en = tk.Entry(stage_page, textvariable=var_stageID, width=11)
    stageID_en.place(x=140 , y=40)

    stageName_La = tk.Label(stage_page, text='Stage Name:')
    stageName_La.place(x=50, y=70)
    stageName_en = tk.Entry(stage_page, textvariable=var_stageName, width=35)
    stageName_en.place(x=140, y=70)

    stageDateLa = tk.Label(stage_page, text='Date:')
    stageDateLa.place(x=50, y=100)

    stageDateLa_y = tk.Label(stage_page, text='Y:')
    stageDateLa_y.place(x=130, y=100)

    stage_startDate_y = tk.Entry(stage_page, textvariable=var_stageStartDate_y, width=5)
    stage_startDate_y.place(x=150, y=100)

    stageDateLa_m = tk.Label(stage_page, text='M:')
    stageDateLa_m.place(x=190, y=100)

    stage_startDate_m = tk.Entry(stage_page, textvariable=var_stageStartDate_m, width=5)
    stage_startDate_m.place(x=215, y=100)

    to = tk.Label(stage_page, text='—')
    to.place(x=260, y=100)

    stageDateLa_toy = tk.Label(stage_page, text='Y:')
    stageDateLa_toy.place(x=280, y=100)

    stage_endDate_y = tk.Entry(stage_page, textvariable=var_stage_endDate_y, width=5)
    stage_endDate_y.place(x=300, y=100)

    stageDateLa_tom = tk.Label(stage_page, text='M:')
    stageDateLa_tom.place(x=340, y=100)

    stage_endDate_m = tk.Entry(stage_page, textvariable=var_stage_endDate_m, width=5)
    stage_endDate_m.place(x=365, y=100)

    addstageDate = tk.Button(stage_page, text='添 加', width='5', command=addstageDate)
    updateStageDate = tk.Button(stage_page, text='修改', width='5', command=updateStageDate)
    removeStageDate = tk.Button(stage_page, text='删除', width='5', command=removeStageDate)
    addstageDate.place(x=200, y=190)
    updateStageDate.place(x=255, y=190)
    removeStageDate.place(x=310, y=190)


    #stage列表
    stage_listbox = tk.Listbox(window, height=8, width=24,yscrollcommand=scrolly.set)
    scrolly.config(command=stage_listbox.yview)
    stageData_Z = {}
    stageDate = open(stageDataFile, 'r')
    content_stage = stageDate.readlines()
    for i in content_stage:
        stage = eval(i)
        stageData_Z = stage
        for u in stage.values():
            stage_listbox.insert('end', u[0][5:]+'   '+'('+u[2]+','+u[3]+'-'+u[4]+','+u[5]+')')
    stage_listbox.place(x=530, y=166)
    stage_listbox.bind('<Button-1>', stage1)


    # 排序_升序
    def callBack(value):
        date1 = []
        t = tree.get_children()
        for i in t:
            date1.append(tree.item(i, 'values'))
        items = tree.get_children()
        [tree.delete(item) for item in items]
        ss = {}
        content = []
        for s in date1:
            date2 = []
            date2.append(s[0])
            date2.append(s[1])
            date2.append(s[2])
            date2.append(int(s[3]))
            date2.append(s[4])
            date2.append(int(s[5]))
            date2.append(float(s[6]))
            ss[s[0]] = date2
        for i in ss.values():
            content.append(i)
        content.sort(key=operator.itemgetter(value), reverse=True)
        for s in content:
            tree.insert('', 0, values=(s[0], s[1], s[2], s[3], s[4], s[5], s[6]))

    # 排序_降序
    def callBack_order(value):
        date1 = []
        t = tree.get_children()
        for i in t:
            date1.append(tree.item(i, 'values'))
        items = tree.get_children()
        [tree.delete(item) for item in items]
        ss = {}
        content = []
        for s in date1:
            date2 = []
            date2.append(s[0])
            date2.append(s[1])
            date2.append(s[2])
            date2.append(int(s[3]))
            date2.append(s[4])
            date2.append(int(s[5]))
            date2.append(float(s[6]))
            ss[s[0]] = date2
        for i in ss.values():
            content.append(i)
        content.sort(key=operator.itemgetter(value))
        for s in content:
            tree.insert('', 0, values=(s[0], s[1], s[2], s[3], s[4], s[5], s[6]))


    #判断字符串内容是否为数字
    def check(a):
        if type(a) is not str:
            return False
        else:
            for i in a:
                if i not in string.digits:
                    return False
            return True
    raise_frame(bills_page)

    #填写pdf标题和id
    def pdfWindow():
        def confirmpdf():
            if var_pdftitle.get() != '' and var_pdfid.get() != '':
                educePDF()
                var_pdftitle.set('')
                var_pdfid.set('')
                PDFwindow.destroy()
            else:
                tk.messagebox.showinfo(title='提示', message='请填写内容')
        def cancelpdf():
            var_pdftitle.set('')
            var_pdfid.set('')
        PDFwindow = tk.Toplevel(window)
        PDFwindow.title('xxx律师所')
        PDFwindow.geometry('500x300')
        PDFwindow.maxsize(500, 300)
        PDFwindow.minsize(500, 300)
        hintLa=tk.Label(PDFwindow,text='请填写导出PDF文件标题和编号').place(x=150, y=25)
        pdftitle_label = tk.Label(PDFwindow, text='标题:').place(x=100, y=100)
        pdfid_label = tk.Label(PDFwindow, text='编号:').place(x=100, y=140)
        pdftitle_entry = tk.Entry(PDFwindow, textvariable=var_pdftitle, width=35)
        pdftitle_entry.place(x=150, y=100)
        pdfid_entry = tk.Entry(PDFwindow, textvariable=var_pdfid, width=35)
        pdfid_entry.place(x=150, y=140)
        confirm = tk.Button(PDFwindow, width=6, text='确定', command=confirmpdf)
        confirm .place(x=210, y=202)
        cancel = tk.Button(PDFwindow, width=6, text='清空', command=cancelpdf)
        cancel .place(x=280, y=202)
        PDFwindow.mainloop()

    # 填写DOCX标题和id
    def docxWindow():
        def confirmdocx():
            if var_docxtitle.get() != '' and var_docxid.get() != '':
                plot()
                var_docxtitle.set('')
                var_docxid.set('')
                docxwindow.destroy()
            else:
                tk.messagebox.showinfo(title='提示', message='请填写内容')
        def canceldocx():
            var_docxtitle.set('')
            var_docxid.set('')

        docxwindow = tk.Toplevel(window)
        docxwindow.title('xxx律师所')
        docxwindow.geometry('500x300')
        docxwindow.maxsize(500, 300)
        docxwindow.minsize(500, 300)
        hintLa = tk.Label(docxwindow, text='请填写导出DOCX文件标题和编号').place(x=150, y=25)
        pdftitle_label = tk.Label(docxwindow, text='标题:').place(x=100, y=100)
        pdfid_label = tk.Label(docxwindow, text='编号:').place(x=100, y=140)
        pdftitle_entry = tk.Entry(docxwindow, textvariable=var_docxtitle, width=35)
        pdftitle_entry.place(x=150, y=100)
        pdfid_entry = tk.Entry(docxwindow, textvariable=var_docxid, width=35)
        pdfid_entry.place(x=150, y=140)
        confirm = tk.Button(docxwindow, width=6, text='确定', command=confirmdocx)
        confirm.place(x=210, y=202)
        cancel = tk.Button(docxwindow, width=6, text='清空', command=canceldocx)
        cancel.place(x=280, y=202)
        docxwindow.mainloop()


    window.mainloop()